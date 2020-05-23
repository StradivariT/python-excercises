import json
import docx
from reportlab.pdfgen import canvas

from user import User

class CRUD:
    def __init__(self, users: list):
        self.dbPath = "./Proyecto/database/db.json"
        self.users = users

    def addNewUser(self, name: str, email: str, phone: str, address: str):
        self.users.append(User.fromForm(name, email, phone, address))
        self.saveUsersToDB()

    def modifyUser(self, index: int, name: str, email: str, phone: str, address: str) -> str:
        if (User.isEmpty()):
            return "No users to modify"

        if (index < 0 or index >= len(self.users)):
            return "Invalid index"

        self.users[index].modify(name, email, phone, address)
        self.saveUsersToDB()

        return ""

    def deleteUser(self, index: int) -> str:
        if (User.isEmpty()):
            return "No users to delete"

        if (index < 0 or index >= len(self.users)):
            return "Invalid index"

        del self.users[index]
        self.saveUsersToDB()

        return ""

    def saveUsersToDB(self):
        jsonStr = json.dumps([user.__dict__ for user in self.users])
        
        db = open(self.dbPath, "w")
        db.write(jsonStr)
        db.close()

    def generateFiles(self, userIndexes: str, allUsers: bool, plainText: bool, pdf: bool, word: bool) -> str:
        if (User.isEmpty()):
            return "No users to generate files"

        funcs = []
        if (plainText): 
            funcs.append(self.generatePlainText)
        
        if (pdf): 
            funcs.append(self.generatePDF)

        if (word): 
            funcs.append(self.generateWord)

        users = []
        if (allUsers):
            users = self.users
        else:
            indexes = list(map(lambda x : int(x), userIndexes.split(",")))
            users = list(map(lambda x : self.users[x], indexes))

        list(map(lambda x : list(map(lambda y : y(x), funcs)), users))
        return ""
    
    def generatePlainText(self, user: User):
        plainFile = open(f"./Proyecto/output/plain/{ user.name }_{ user.id }", "w")
        plainFile.write(user.formatFile())
        plainFile.close()

    def generatePDF(self, user: User):
        pdf = canvas.Canvas(f"./Proyecto/output/pdf/{ user.name }_{ user.id }.pdf")
        pdf.setTitle(f"{ user.name } information")
        
        pdf.setFont("Courier", 30)
        pdf.drawCentredString(300, 770, f"User: { user.name }")

        pdf.line(100, 750, 500, 750)

        pdf.setFont("Courier", 20)
        
        pdf.drawString(100, 690, f"ID: { user.id }")
        pdf.drawString(100, 660, f"Email: { user.email }")
        pdf.drawString(100, 630, f"Phone: { user.phone }")
        pdf.drawString(100, 600, f"Address: { user.address }")

        pdf.save()

    def generateWord(self, user: User):
        doc = docx.Document()

        doc.add_paragraph(f"ID: { user.id }")
        doc.add_paragraph(f"User: { user.name }")
        doc.add_paragraph(f"Email: { user.email }")
        doc.add_paragraph(f"Phone: { user.phone }")
        doc.add_paragraph(f"Address: { user.address }")

        doc.save(f"./Proyecto/output/word/{ user.name }_{ user.id }.docx")