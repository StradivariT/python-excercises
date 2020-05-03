import json
import docx
from reportlab.pdfgen import canvas

import utils as u
from user import User

def getFileOption() -> int:
    print("Pick which output file format to generate:")
    print("1. Plain Text")
    print("2. PDF")
    print("3. Microsoft Word")

    return int(input())

def generatePlainText(user: User):
    plainFile = open(f"./Proyecto/output/plain/{ user.name }_{ user.id }", "w")
    plainFile.write(str(user))
    plainFile.close()

def generatePDF(user: User):
    pdf = canvas.Canvas(f"./Proyecto/output/pdf/{ user.name }_{ user.id }.pdf")
    pdf.setTitle(f"{ user.name } information")
    
    pdf.setFont("Courier", 30)
    pdf.drawCentredString(300, 770, f"User: { user.name }")

    pdf.line(100, 750, 500, 750)

    pdf.setFont("Courier", 20)
    
    pdf.drawString(100, 690, f"Email: { user.email }")
    pdf.drawString(100, 660, f"Phone: { user.phone }")
    pdf.drawString(100, 630, f"Address: { user.address }")

    pdf.save()

def generateWord(user: User):
    doc = docx.Document()

    doc.add_paragraph(f"User: { user.name }")
    doc.add_paragraph(f"Email: { user.email }")
    doc.add_paragraph(f"Phone: { user.phone }")
    doc.add_paragraph(f"Address: { user.address }")

    doc.save(f"./Proyecto/output/word/{ user.name }_{ user.id }.docx")

def getFileGenerationFunction():
    return {
        1: generatePlainText,
        2: generatePDF,
        3: generateWord
    } [getFileOption()]

def generateFileForSingleUser(users: list):
    if User.isEmpty():
        print("There are no users registered")
        return

    option = u.getSingleUserOption("Pick which user you want to generate the file for:", users)
    getFileGenerationFunction()(users[option])

    print("File generated on output folder")

def generateFileForAllUsers(users: list):
    if User.isEmpty():
        print("There are no users registered")
        return

    list(map(getFileGenerationFunction(), users))

    print("Files generated on output folder")

def generateFileForSomeUsers(users: list):
    if User.isEmpty():
        print("There are no users registered")
        return

    options = u.getMultipleUsersOptions("Pick which users you want to generate files for separated by commas, ex. 1,2,...:", users)
    func = getFileGenerationFunction()

    list(map(lambda x: func(users[x]), options))

    print("Files generated on output folder")